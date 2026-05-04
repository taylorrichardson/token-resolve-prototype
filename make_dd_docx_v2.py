import docx

def main():
    template_path = '/Users/taylor.richardson/.gemini/skills/design-document-writer/assets/Design Document Template.docx'
    doc = docx.Document(template_path)

    # Replace Title and Subtitle
    for p in doc.paragraphs:
        if p.style.name == 'Title' and 'Feature Name' in p.text:
            p.text = 'Automated Token Resolution'
        elif p.style.name == 'Subtitle':
            p.text = 'Real-time, reactive token resolution system for Vault Content Plans'

    # Remove all paragraphs that are NOT headings or TOC, but keep them if they are the "Release Note" text.
    paragraphs_to_delete = []
    for p in doc.paragraphs:
        if p.style.name.startswith('Heading') or p.style.name in ['Title', 'Subtitle'] or p.style.name.startswith('toc'):
            continue
        if p.text.strip() == 'Release Note':
            continue
        
        paragraphs_to_delete.append(p)
        
    for p in paragraphs_to_delete:
        p._element.getparent().remove(p._element)

    content_map = {
        "Feature Info": ["VPD Feature: [Link Placeholder]", "Jira EPIC: [Link Placeholder]"],
        "Required Reading": ["Feature Overview: [Link Placeholder to FO]", "UI Flow: N/A"],
        "Release Note": ["Replaces the manual token resolution process with a fully automated, real-time mechanism."],
        "Background & Stakeholders": [
            "Currently, token resolution in Vault Content Plans requires users to manually execute the \"Update Tokens in Fields\" or \"Update Tokens in Field\" actions. This manual step often leads to delayed error discovery (such as \"publishing string too long\" errors) and prevents users from immediately visualizing how their final published output will be structured.",
            "This feature introduces a real-time, reactive token resolution system, replacing the manual process. It aims to reduce customer frustration and publishing errors by allowing users to preview and resolve tokens instantly via UI interactions such as drag-and-drop document matching.",
            "Key Personas:",
            "Submission Publisher",
            "Content Plan Author"
        ],
        "Competitor Analysis": ["N/A"],
        "Terminology & Personas": ["N/A"],
        "Use Cases": ["N/A"],
        "Data Model Impacts": [
            "To support non-destructive manual overrides without losing the original template token strings, a new field will be introduced.",
            "Object: Content Plan Item (content_plan_item__v)",
            "New Field: manual_override__v (Text)",
            "Purpose: Stores the user's manual override of the resolved string when a document is attached."
        ],
        "Standard Layouts": [
            "The manual_override__v field shall be added to the standard Content Plan Item page layouts, though primarily interacted with via the inline grid edit."
        ],
        "Data Model - Vault Object Configuration": ["N/A"],
        "Dependencies": ["N/A"],
        "Feature Requirements": [
            "5.1 Real-Time Resolution",
            "The system shall continuously evaluate and resolve tokens based on the current metadata and matched documents.",
            "(P1) The system shall automatically resolve tokens (e.g., ${matched_document.name__v}) in real-time as metadata changes or documents are matched to Content Plan Items.",
            "(P1) The manual \"Update Tokens in Fields\" and \"Update Tokens in Field\" user actions shall be deprecated and removed from the UI.",
            "",
            "5.2 Folder Structure View & Drag-and-Drop Matching",
            "The UI will be updated to focus on the final deliverable's hierarchy.",
            "(P1) The system shall provide a dedicated folder structure view replacing the standard lifecycle states view to dynamically visualize the post-publishing layout.",
            "(P1) The UI shall allow users to drag and drop documents directly from a sidebar onto Content Plan Items in the grid.",
            "(P1) When a document is dropped onto an item, the system shall instantly simulate document matching and update the UI to display the resolved token values.",
            "",
            "5.3 Live Path Normalization",
            "The Published Output Location (POL) generation must be transparent and instantaneous.",
            "(P1) The system shall automatically generate and normalize the Published Output Location (published_output_location__v) immediately in the UI.",
            "(P1) Normalization rules shall include removing special characters (~ \\ : * ? < > \" |), enforcing lowercase, and truncating multiple consecutive hyphens into a single hyphen.",
            "",
            "5.4 Publishing Validation",
            "The system must actively prevent incomplete plans from proceeding to publishing.",
            "(P1) Unresolved tokens shall be explicitly highlighted (e.g., in blue) within the UI to signal missing data to the user.",
            "(P1) The system shall visually gate and disable the \"Move to Publishing\" action until all required Content Plan Items have successfully matched documents attached and no unresolved tokens remain.",
            "",
            "5.5 Non-Destructive Overrides",
            "Users require the ability to make manual adjustments without permanently breaking the link to the underlying template.",
            "(P1) When a user clicks on an item without an attached document, the system shall allow inline editing of the raw template token string.",
            "(P1) When a user clicks on an item with an attached document, the system shall allow inline editing to create a manual override of the resolved string. This override shall be saved to the new manual_override__v field.",
            "(P1) If an attached document is removed, the system shall clear the manual_override__v value and revert the displayed name to the original template token string."
        ],
        "Accessibility Requirements": ["N/A"],
        "Upgrade Requirements": ["N/A"],
        "Feature Usage Tracking": ["N/A"],
        "App Stats": ["N/A"],
        "Business Activity Logging": ["N/A"],
        "Customer Activity Logging": ["N/A"],
        "Future Considerations & Roadmap": ["N/A"],
        "Vault Connections": ["N/A"],
        "Performance Considerations and Limits": [
            "Token resolution and path normalization must execute in under 1 second to ensure a responsive UI."
        ],
        "Security Considerations": [
            "Field-level security for manual_override__v shall mirror existing edit permissions for the Content Plan Item name/template fields."
        ],
        "Audit Trail Requirements": ["N/A"],
        "Document Audit Trail": ["N/A"],
        "Record Audit Trail": [
            "Changes to the manual_override__v field shall be tracked in the standard Vault audit trail for the Content Plan Item."
        ],
        "System Audit Log": ["N/A"],
        "MDL Requirements": ["N/A"],
        "Testing Considerations": [
            "Verify real-time resolution upon drag-and-drop matching.",
            "Verify live path normalization accurately removes all restricted characters and forces lowercase.",
            "Verify the \"Move to Publishing\" action is strictly disabled when unresolved tokens are present.",
            "Verify that clearing a matched document properly clears the manual_override__v field and restores the template token."
        ],
        "Discussions & Decisions": ["N/A"]
    }

    # Find paragraphs before inserting new ones so indices don't shift during search
    heading_paras = {}
    for h_text in content_map.keys():
        for p in doc.paragraphs:
            if h_text in p.text and (p.style.name.startswith('Heading') or h_text == 'Release Note'):
                heading_paras[h_text] = p
                break

    for h_text, lines in content_map.items():
        if h_text in heading_paras:
            target_p = heading_paras[h_text]
            idx = doc.paragraphs.index(target_p)
            
            # Insert logic: we insert in reverse order BEFORE the next paragraph to maintain sequence
            next_p = doc.paragraphs[idx+1] if idx + 1 < len(doc.paragraphs) else None
            
            for line in reversed(lines):
                if line == "": continue # skip empty lines
                style = "Normal"
                if line.startswith("(P1)") or line.startswith("Object:") or line.startswith("New Field:") or line.startswith("Purpose:") or line.startswith("Verify") or line in ["Submission Publisher", "Content Plan Author"]:
                    style = "List Paragraph"
                    
                if next_p:
                    next_p.insert_paragraph_before(line, style=style)
                else:
                    doc.add_paragraph(line, style=style)

    doc.save('/Users/taylor.richardson/Projects/prototypes/Token-Resolve/token-resolve-app/documentation/DESIGN_DOCUMENT.docx')
    print("Done")

if __name__ == '__main__':
    main()